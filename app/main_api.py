# --- KAGGLE-POWERED RAG SYSTEM - COMPLETE 1144+ LINES WITH DEADLOCK FIX ---
#final version 
import os
import json
import uuid
import time
import re
import asyncio
import logging
import hashlib
import httpx
from typing import List, Dict, Any, Optional
from collections import defaultdict
from itertools import cycle
from pathlib import Path
import functools
import threading
import concurrent.futures

# FastAPI and core dependencies
from fastapi import FastAPI, Body, HTTPException, Request, Depends, Header
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# LangChain imports
from langchain_community.vectorstores import Chroma

# Multi-format document processing
import fitz  # PyMuPDF
import pdfplumber
import docx
import openpyxl
import csv
import zipfile
import email
from email.policy import default
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET

# LLM providers
import groq
import openai
import google.generativeai as genai

import cachetools
from dotenv import load_dotenv

# Setup
load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Kaggle-Powered Hackathon RAG", version="5.4.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*", "ngrok-skip-browser-warning"],
)

# --- CRITICAL FIX: LAZY KAGGLE MODEL CLIENT ---
class LazyKaggleModelClient:
    """LAZY INITIALIZATION: Only connects when actually needed - PREVENTS 'Preparing Space' ISSUE"""
    def __init__(self):
        self._client = None
        self._endpoint = None
        self._initialized = False
        logger.info("🎯 Lazy Kaggle Model Client created (no immediate connection)")
    
    def _initialize_if_needed(self):
        """Initialize client only when first API call is made"""
        if not self._initialized:
            # Get endpoint from Hugging Face Secrets (or fallback to env var)
            self._endpoint = os.getenv("KAGGLE_NGROK_URL") or os.getenv("KAGGLE_ENDPOINT", "")
            
            if not self._endpoint:
                logger.error("❌ No KAGGLE_NGROK_URL found in secrets or environment!")
                raise Exception("Kaggle endpoint not configured")
            
            self._endpoint = self._endpoint.rstrip('/')
            self._client = httpx.AsyncClient(
                timeout=30.0,
                headers={"ngrok-skip-browser-warning": "true"}
            )
            self._initialized = True
            logger.info(f"🎯 Lazy Kaggle client initialized: {self._endpoint}")
    
    async def health_check(self) -> bool:
        """Check if Kaggle model server is healthy"""
        try:
            self._initialize_if_needed()
            response = await self._client.get(f"{self._endpoint}/health")
            return response.status_code == 200
        except Exception as e:
            logger.error(f"Kaggle health check failed: {e}")
            return False
    
    async def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings using Kaggle GPU"""
        try:
            self._initialize_if_needed()
            response = await self._client.post(
                f"{self._endpoint}/embed",
                json={"texts": texts}
            )
            response.raise_for_status()
            result = response.json()
            logger.info(f"🎯 Kaggle embeddings: {result.get('count', 0)} texts in {result.get('processing_time', 0):.2f}s")
            return result["embeddings"]
        except Exception as e:
            logger.error(f"Kaggle embedding error: {e}")
            return []
    
    async def rerank_documents(self, query: str, documents: List[str], k: int = 8) -> List[str]:
        """Rerank documents using Kaggle GPU"""
        try:
            self._initialize_if_needed()
            response = await self._client.post(
                f"{self._endpoint}/rerank",
                json={
                    "query": query,
                    "documents": documents,
                    "k": k
                }
            )
            response.raise_for_status()
            result = response.json()
            logger.info(f"🎯 Kaggle reranking: {k} docs in {result.get('processing_time', 0):.2f}s")
            return result["reranked_documents"]
        except Exception as e:
            logger.error(f"Kaggle reranking error: {e}")
            return documents[:k]

# --- LIGHTWEIGHT QUERY PROCESSOR (YOUR COMPLETE ORIGINAL) ---
class LightweightQueryProcessor:
    def __init__(self, kaggle_client: LazyKaggleModelClient):
        self.kaggle_client = kaggle_client
        self.cache = cachetools.TTLCache(maxsize=500, ttl=3600)
        
    async def enhance_query_semantically(self, question: str, domain: str = "insurance") -> str:
        """OPTIMIZED semantic query processing"""
        
        # Quick cache check with shorter hash
        cache_key = hashlib.md5(question.encode()).hexdigest()[:8]
        if cache_key in self.cache:
            return self.cache[cache_key]
        
        # Streamlined domain expansion
        enhanced_query = self._expand_with_domain_knowledge_fast(question, domain)
        enhanced_query = self._handle_incomplete_questions(enhanced_query)
        
        # Cache result
        self.cache[cache_key] = enhanced_query
        return enhanced_query
    
    def _expand_with_domain_knowledge_fast(self, query: str, domain: str) -> str:
        """OPTIMIZED domain expansion - same intelligence, faster processing"""
        
        # Streamlined expansion mapping for speed
        key_expansions = {
            'grace period': 'payment deadline premium due',
            'waiting period': 'exclusion time coverage delay',
            'pre-existing': 'prior medical condition',
            'coverage': 'policy benefits protection',
            'exclusion': 'limitations restrictions',
            'premium': 'insurance cost payment',
            'claim': 'benefit request reimbursement',
            'ayush': 'alternative medicine treatment',
            'hospital': 'healthcare facility medical center'
        }
        
        query_lower = query.lower()
        for key_term, expansion in key_expansions.items():
            if key_term in query_lower:
                return f"{query}. Also: {expansion}"
        
        return query
    
    def _handle_incomplete_questions(self, query: str) -> str:
        """Handle R4's 'half questions' requirement"""
        incomplete_patterns = [
            r'^(what|how|when|where|why)\s*\?*$',
            r'^(yes|no)\s*\?*$',
            r'^\w{1,3}\s*\?*$',
            r'^(this|that|it)\s*',
        ]
        
        query_lower = query.lower()
        is_incomplete = any(re.search(pattern, query_lower) for pattern in incomplete_patterns)
        
        if is_incomplete and len(query.split()) <= 2:
            return f"{query}. Please provide information about insurance policy terms, coverage, exclusions, waiting periods, or benefits."
        
        return query

# --- ANTI-JAILBREAK SECURITY SYSTEM (YOUR COMPLETE ORIGINAL) ---
class SecurityGuard:
    def __init__(self):
        self.jailbreak_patterns = [
            r'ignore.*previous.*instructions',
            r'act.*as.*different.*character',
            r'generate.*code.*(?:javascript|python|html)',
            r'write.*program',
            r'roleplay.*as',
            r'pretend.*you.*are',
            r'system.*prompt',
            r'override.*settings',
            r'bypass.*restrictions',
            r'admin.*mode',
            r'developer.*mode',
            r'tell.*me.*about.*yourself',
            r'what.*are.*you',
            r'who.*created.*you'
        ]
    
    def detect_jailbreak(self, text: str) -> bool:
        """Detect jailbreak attempts"""
        text_lower = text.lower()
        return any(re.search(pattern, text_lower) for pattern in self.jailbreak_patterns)
    
    def sanitize_response(self, question: str, answer: str) -> str:
        """Sanitize responses against jailbreaks"""
        if self.detect_jailbreak(question):
            return "I can only provide information based on the document content provided. Please ask questions about the document."
        
        # Remove any potential code or script tags
        answer = re.sub(r'<script.*?</script>', '', answer, flags=re.DOTALL | re.IGNORECASE)
        answer = re.sub(r'<.*?>', '', answer)  # Remove HTML tags
        
        return answer

# --- MULTI-LLM MANAGER (YOUR COMPLETE ORIGINAL WITH ALL PROVIDERS) ---
class MultiLLMManager:
    def __init__(self):
        # Initialize multiple LLM providers with fallback
        self.providers = ['groq']  # Start with Groq as primary
        
        self.groq_keys = cycle([k.strip() for k in os.getenv("GROQ_API_KEYS", "").split(',') if k.strip()])
        
        # Optional paid providers (if keys available)
        openai_keys = [k.strip() for k in os.getenv("OPENAI_API_KEYS", "").split(',') if k.strip()]
        gemini_keys = [k.strip() for k in os.getenv("GEMINI_API_KEYS", "").split(',') if k.strip()]
        
        if openai_keys:
            self.providers.append('openai')
            self.openai_keys = cycle(openai_keys)
            
        if gemini_keys:
            self.providers.append('gemini') 
            self.gemini_keys = cycle(gemini_keys)
        
        self.current_provider_index = 0
        logger.info(f"🔑 Multi-LLM Manager initialized with {len(self.providers)} providers")
    
    async def get_response(self, prompt: str, max_tokens: int = 900) -> str:
        """Get response with automatic fallback between providers"""
        for attempt in range(len(self.providers)):
            try:
                provider = self.providers[self.current_provider_index]
                
                if provider == 'groq':
                    return await self._groq_response(prompt, max_tokens)
                elif provider == 'openai':
                    return await self._openai_response(prompt, max_tokens)
                elif provider == 'gemini':
                    return await self._gemini_response(prompt, max_tokens)
                    
            except Exception as e:
                logger.warning(f"{provider} failed: {e}")
                self.current_provider_index = (self.current_provider_index + 1) % len(self.providers)
                continue
        
        return "Error: All LLM providers failed"
    
    async def _groq_response(self, prompt: str, max_tokens: int) -> str:
        key = next(self.groq_keys)
        client = groq.Groq(api_key=key)
        
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=max_tokens,
            top_p=0.9
        )
        return response.choices[0].message.content.strip()
    
    async def _openai_response(self, prompt: str, max_tokens: int) -> str:
        key = next(self.openai_keys)
        openai.api_key = key
        
        response = await openai.ChatCompletion.acreate(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=max_tokens
        )
        return response.choices[0].message.content.strip()
    
    async def _gemini_response(self, prompt: str, max_tokens: int) -> str:
        key = next(self.gemini_keys)
        genai.configure(api_key=key)
        
        model = genai.GenerativeModel('gemini-pro')
        response = await model.generate_content_async(prompt)
        return response.text.strip()

# --- COMPLETE UNIVERSAL DOCUMENT PROCESSOR (ALL YOUR ORIGINAL FEATURES) ---
class UniversalDocumentProcessor:
    def __init__(self):
        # SPEED OPTIMIZATIONS: Reduced limits
        self.chunk_size = 1000      # Reduced from 1200
        self.chunk_overlap = 200
        self.max_chunks = 200       # Kept at 200 (good balance)
        self.max_pages = 18         # Reduced from 25
        
        # Smaller cache for speed
        self.cache = cachetools.TTLCache(maxsize=50, ttl=1800)
        
        # Supported formats (KEEPING all your excellent processors)
        self.processors = {
            '.pdf': self.process_pdf,
            '.docx': self.process_docx,
            '.doc': self.process_doc,
            '.xlsx': self.process_excel,
            '.xls': self.process_excel,
            '.csv': self.process_csv,
            '.txt': self.process_text,
            '.html': self.process_html,
            '.xml': self.process_xml,
            '.eml': self.process_email,
            '.zip': self.process_archive,
            '.json': self.process_json
        }
        
        logger.info("⚡ Speed-Optimized Universal Document Processor initialized")
    
    def get_file_hash(self, content: bytes) -> str:
        """Generate shorter hash for caching"""
        return hashlib.md5(content).hexdigest()[:8]
    
    async def process_document(self, file_path: str, content: bytes) -> List[Dict[str, Any]]:
        """Process any document format with optimized caching"""
        file_hash = self.get_file_hash(content)
        
        # Check cache first
        if file_hash in self.cache:
            logger.info(f"📦 Cache hit for {os.path.basename(file_path)}")
            return self.cache[file_hash]
        
        # Detect file type
        file_ext = Path(file_path).suffix.lower()
        if not file_ext:
            file_ext = self._detect_file_type(content)
        
        # Process based on file type
        processor = self.processors.get(file_ext, self.process_text)
        
        try:
            chunks = await processor(file_path, content)
            
            # Cache the result
            self.cache[file_hash] = chunks
            
            logger.info(f"✅ Processed {os.path.basename(file_path)}: {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"❌ Processing failed for {file_path}: {e}")
            return self._emergency_text_extraction(content, file_path)
    
    def _detect_file_type(self, content: bytes) -> str:
        """Detect file type from content"""
        if content.startswith(b'%PDF'):
            return '.pdf'
        elif content.startswith(b'PK'):
            return '.docx' if b'word/' in content[:1000] else '.zip'
        elif content.startswith(b'<html') or content.startswith(b'<!DOCTYPE'):
            return '.html'
        elif content.startswith(b'<?xml'):
            return '.xml'
        else:
            return '.txt'
    
    # --- SPEED-OPTIMIZED PDF PROCESSING (YOUR COMPLETE ORIGINAL) ---
    async def process_pdf(self, file_path: str, content: bytes) -> List[Dict[str, Any]]:
        """Enhanced PDF processing with speed optimizations"""
        chunks = []
        temp_path = f"/tmp/{uuid.uuid4().hex[:6]}.pdf"  # Shorter UUID
        
        with open(temp_path, 'wb') as f:
            f.write(content)
        
        try:
            # Extract text with PyMuPDF
            doc = fitz.open(temp_path)
            full_text = ""
            
            # SPEED OPTIMIZATION: Process fewer pages
            for page_num in range(min(len(doc), self.max_pages)):
                page = doc[page_num]
                text = page.get_text()
                
                if text.strip():
                    full_text += f"\n\nPage {page_num + 1}:\n{self._clean_text(text)}"
            
            doc.close()
            
            # OPTIMIZED table extraction
            table_text = await self._extract_pdf_tables_fast(temp_path)
            if table_text:
                full_text += f"\n\n=== TABLES ===\n{table_text}"
            
            # Create semantic chunks
            chunks = self._create_semantic_chunks(full_text, file_path, "pdf")
            
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            chunks = self._emergency_text_extraction(content, file_path)
        
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
        return chunks
    
    async def _extract_pdf_tables_fast(self, file_path: str) -> str:
        """SPEED-OPTIMIZED table extraction"""
        table_text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                # SPEED OPTIMIZATION: Fewer pages and tables
                for page_num, page in enumerate(pdf.pages[:10]):  # Reduced from 12
                    tables = page.find_tables()
                    for i, table in enumerate(tables[:1]):  # Only 1 table per page
                        try:
                            table_data = table.extract()
                            if table_data and len(table_data) > 1:
                                table_md = f"\n**Table {i+1} (Page {page_num+1})**\n"
                                for row in table_data[:12]:  # Reduced from 15
                                    if row:
                                        clean_row = [str(cell or "").strip()[:30] for cell in row]
                                        table_md += "| " + " | ".join(clean_row) + " |\n"
                                table_text += table_md + "\n"
                        except:
                            continue
        except Exception as e:
            logger.warning(f"Table extraction failed: {e}")
        
        return table_text
    
    # --- OTHER FORMAT PROCESSORS (ALL YOUR EXCELLENT FEATURES) ---
    async def process_docx(self, file_path: str, content: bytes) -> List[Dict[str, Any]]:
        """Process DOCX files"""
        temp_path = f"/tmp/{uuid.uuid4().hex[:6]}.docx"
        with open(temp_path, 'wb') as f:
            f.write(content)
        
        try:
            doc = docx.Document(temp_path)
            full_text = ""
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                table_text = "\n**TABLE**\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text += "| " + " | ".join(row_text) + " |\n"
                full_text += table_text + "\n"
            
            chunks = self._create_semantic_chunks(full_text, file_path, "docx")
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            chunks = self._emergency_text_extraction(content, file_path)
        
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
        return chunks
    
    async def process_doc(self, file_path: str, content: bytes) -> List[Dict[str, Any]]:
        """Process DOC files (fallback to text extraction)"""
        return self._emergency_text_extraction(content, file_path)
    
    async def process_excel(self, file_path: str, content: bytes) -> List[Dict[str, Any]]:
        """Process Excel files"""
        temp_path = f"/tmp/{uuid.uuid4().hex[:6]}.xlsx"
        with open(temp_path, 'wb') as f:
            f.write(content)
        
        try:
            workbook = openpyxl.load_workbook(temp_path, read_only=True)
            full_text = ""
            
            for sheet_name in workbook.sheetnames[:3]:
                sheet = workbook[sheet_name]
                full_text += f"\n**Sheet: {sheet_name}**\n"
                
                for row_num, row in enumerate(sheet.iter_rows(max_row=50, values_only=True)):
                    if row_num == 0 or any(cell for cell in row):
                        row_text = [str(cell or "").strip()[:30] for cell in row[:8]]
                        full_text += "| " + " | ".join(row_text) + " |\n"
            
            workbook.close()
            chunks = self._create_semantic_chunks(full_text, file_path, "excel")
            
        except Exception as e:
            logger.error(f"Excel processing error: {e}")
            chunks = self._emergency_text_extraction(content, file_path)
        
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
        return chunks
    
    # --- Other format processors (keeping all your excellent features) ---
    async def process_csv(self, file_path: str, content: bytes) -> List[Dict[str, Any]]:
        try:
            text_content = content.decode('utf-8', errors='ignore')
            lines = text_content.split('\n')
            
            full_text = "**CSV DATA**\n"
            for i, line in enumerate(lines[:100]):
                if line.strip():
                    full_text += f"| {line} |\n"
            
            return self._create_semantic_chunks(full_text, file_path, "csv")
        except Exception as e:
            logger.error(f"CSV processing error: {e}")
            return []
    
    async def process_text(self, file_path: str, content: bytes) -> List[Dict[str, Any]]:
        try:
            text = content.decode('utf-8', errors='ignore')
            return self._create_semantic_chunks(text, file_path, "text")
        except Exception as e:
            logger.error(f"Text processing error: {e}")
            return []
    
    async def process_html(self, file_path: str, content: bytes) -> List[Dict[str, Any]]:
        try:
            soup = BeautifulSoup(content, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text()
            return self._create_semantic_chunks(text, file_path, "html")
        except Exception as e:
            logger.error(f"HTML processing error: {e}")
            return []
    
    async def process_xml(self, file_path: str, content: bytes) -> List[Dict[str, Any]]:
        try:
            root = ET.fromstring(content)
            def extract_text(element, level=0):
                text = ""
                if element.text and element.text.strip():
                    text += f"{'  ' * level}{element.tag}: {element.text.strip()}\n"
                for child in element:
                    text += extract_text(child, level + 1)
                return text
            full_text = extract_text(root)
            return self._create_semantic_chunks(full_text, file_path, "xml")
        except Exception as e:
            logger.error(f"XML processing error: {e}")
            return []
    
    async def process_email(self, file_path: str, content: bytes) -> List[Dict[str, Any]]:
        try:
            msg = email.message_from_bytes(content, policy=default)
            full_text = f"**EMAIL**\n"
            full_text += f"From: {msg.get('From', 'Unknown')}\n"
            full_text += f"Subject: {msg.get('Subject', 'No Subject')}\n\n"
            
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body = part.get_content()
                        full_text += f"Content:\n{body}\n"
            else:
                body = msg.get_content()
                full_text += f"Content:\n{body}\n"
            
            return self._create_semantic_chunks(full_text, file_path, "email")
        except Exception as e:
            logger.error(f"Email processing error: {e}")
            return []
    
    async def process_archive(self, file_path: str, content: bytes) -> List[Dict[str, Any]]:
        temp_path = f"/tmp/{uuid.uuid4().hex[:6]}.zip"
        with open(temp_path, 'wb') as f:
            f.write(content)
        
        chunks = []
        try:
            if file_path.endswith('.zip'):
                with zipfile.ZipFile(temp_path, 'r') as zip_file:
                    for file_info in zip_file.filelist[:5]:
                        try:
                            file_content = zip_file.read(file_info)
                            sub_chunks = await self.process_document(file_info.filename, file_content)
                            chunks.extend(sub_chunks[:15])  # Limit sub-chunks for speed
                        except:
                            continue
        except Exception as e:
            logger.error(f"Archive processing error: {e}")
        
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
        return chunks
    
    async def process_json(self, file_path: str, content: bytes) -> List[Dict[str, Any]]:
        try:
            data = json.loads(content.decode('utf-8'))
            full_text = json.dumps(data, indent=2, ensure_ascii=False)
            return self._create_semantic_chunks(full_text, file_path, "json")
        except Exception as e:
            logger.error(f"JSON processing error: {e}")
            return []
    
    # --- UTILITY METHODS (YOUR EXCELLENT ORIGINAL) ---
    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'\s+', ' ', text)
        
        # Remove noise patterns
        noise_patterns = [
            r'Office of.*Insurance Ombudsman.*?\n',
            r'Lalit Bhawan.*?\n',
            r'^\d+\s*$'
        ]
        
        for pattern in noise_patterns:
            text = re.sub(pattern, '', text, flags=re.MULTILINE)
        
        return text.strip()
    
    def _create_semantic_chunks(self, text: str, source: str, doc_type: str) -> List[Dict[str, Any]]:
        """Create semantic chunks from text"""
        text = self._clean_text(text)
        
        if not text or len(text) < 50:
            return []
        
        # Smart sentence-based chunking
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= self.chunk_size:
                current_chunk += sentence + " "
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + " "
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        # Convert to structured chunks
        structured_chunks = []
        for i, chunk_text in enumerate(chunks[:self.max_chunks]):
            structured_chunks.append({
                "content": chunk_text,
                "metadata": {
                    "source": os.path.basename(source),
                    "chunk_index": i,
                    "document_type": doc_type,
                    "chunk_length": len(chunk_text)
                },
                "chunk_id": str(uuid.uuid4())
            })
        
        return structured_chunks
    
    def _emergency_text_extraction(self, content: bytes, file_path: str) -> List[Dict[str, Any]]:
        """Emergency text extraction for unsupported formats"""
        try:
            text = content.decode('utf-8', errors='ignore')
            if len(text) > 50:
                return self._create_semantic_chunks(text, file_path, "unknown")
        except:
            pass
        
        return [{
            "content": "Failed to extract content from document",
            "metadata": {
                "source": os.path.basename(file_path),
                "chunk_index": 0,
                "document_type": "error",
                "error": True
            },
            "chunk_id": str(uuid.uuid4())
        }]

# --- GEMINI'S FIX: DEADLOCK-FREE RAG PIPELINE ---
class DeadlockFreeRAGPipeline:
    """FIXED: Direct embedding management - no more AsyncKaggleEmbeddingWrapper deadlock"""
    def __init__(self, collection_name: str, llm_manager: MultiLLMManager, kaggle_client: LazyKaggleModelClient):
        self.collection_name = collection_name
        self.llm_manager = llm_manager
        self.kaggle_client = kaggle_client
        self.security_guard = SecurityGuard()
        self.query_processor = LightweightQueryProcessor(kaggle_client)
        
        # GEMINI'S FIX: No embedding function - let Chroma be a simple data store
        self.vectorstore = Chroma(
            collection_name=collection_name,
            # REMOVED: embedding_function parameter completely
            persist_directory="/tmp/chroma_kaggle"
        )
        
        logger.info(f"🚀 Deadlock-Free RAG Pipeline initialized: {collection_name}")
    
    async def add_documents(self, chunks: List[Dict[str, Any]]):
        """GEMINI'S FIX: Direct embedding management - no deadlock"""
        if not chunks:
            return
        
        logger.info(f"📚 Processing {len(chunks)} chunks...")
        
        # Advanced quality filtering (YOUR EXCELLENT ORIGINAL LOGIC)
        quality_chunks = []
        for chunk in chunks:
            content = chunk['content']
            
            # Skip error chunks
            if chunk['metadata'].get('error'):
                continue
            
            # Quality assessment
            quality_score = 0
            
            # Length factor
            if 100 <= len(content) <= 2000:
                quality_score += 2
            elif len(content) > 50:
                quality_score += 1
            
            # Content richness
            sentences = len(re.split(r'[.!?]+', content))
            if sentences > 3:
                quality_score += 1
            
            # Numerical data (good for policies)
            numbers = len(re.findall(r'\d+', content))
            if numbers > 0:
                quality_score += 1
            
            if quality_score >= 2:
                quality_chunks.append(chunk)
        
        logger.info(f"📚 Filtered to {len(quality_chunks)} quality chunks")
        
        if not quality_chunks:
            return
        
        # GEMINI'S FIX: Step 1 - Get texts
        texts = [chunk['content'] for chunk in quality_chunks[:100]]  # Reduced from 150 for speed
        
        # GEMINI'S FIX: Step 2 - Embed all texts via Kaggle (Manager gets sauce first)
        logger.info(f"🚀 Embedding {len(texts)} chunks via Kaggle...")
        embeddings = await self.kaggle_client.generate_embeddings(texts)
        
        if not embeddings or len(embeddings) != len(texts):
            logger.error("Embedding failed or returned mismatched count.")
            return
        
        # GEMINI'S FIX: Step 3 - Add to Chroma with pre-calculated embeddings
        # This completely avoids the deadlock!
        self.vectorstore.add_texts(
            texts=texts,
            metadatas=[chunk['metadata'] for chunk in quality_chunks[:100]],
            embeddings=embeddings  # Pass vectors directly - no async calls in Chroma!
        )
        
        logger.info(f"✅ Added {len(texts)} documents with embeddings to vector store (DEADLOCK-FREE)")
    
    async def answer_question(self, question: str) -> str:
        """GEMINI'S FIX: Direct query embedding - no deadlock"""
        # Security check
        if self.security_guard.detect_jailbreak(question):
            return self.security_guard.sanitize_response(question, "")
        
        try:
            # Enhanced query processing
            enhanced_question = await self.query_processor.enhance_query_semantically(question)
            
            # GEMINI'S FIX: Step 1 - Embed the query yourself first (Manager gets sauce)
            query_embedding_list = await self.kaggle_client.generate_embeddings([enhanced_question])
            if not query_embedding_list:
                return "I could not process the query for searching."
            
            query_embedding = query_embedding_list[0]
            
            # GEMINI'S FIX: Step 2 - Search using vector directly (no async calls in Chroma)
            relevant_docs = self.vectorstore.similarity_search_by_vector(
                embedding=query_embedding,
                k=15
            )
            
            if not relevant_docs:
                return "I don't have sufficient information to answer this question based on the provided documents."
            
            # Use Kaggle GPU for reranking (GAME CHANGER)
            doc_contents = [doc.page_content for doc in relevant_docs]
            
            if await self.kaggle_client.health_check():
                logger.info("🎯 Using Kaggle GPU for reranking")
                top_docs_content = await self.kaggle_client.rerank_documents(
                    enhanced_question, doc_contents, k=6
                )
            else:
                logger.warning("📦 Kaggle unavailable, using first 6 docs")
                top_docs_content = doc_contents[:6]
            
            # Prepare enhanced context
            context = "\n\n".join(top_docs_content)
            
            # Create advanced semantic prompt
            prompt = self._create_advanced_prompt(context, question)
            
            # Get response from multi-LLM system
            response = await self.llm_manager.get_response(prompt)
            
            # Final security check and cleaning
            response = self.security_guard.sanitize_response(question, response)
            response = self._clean_response(response)
            
            return response
            
        except Exception as e:
            logger.error(f"❌ Question processing failed: {e}")
            return "An error occurred while processing your question."
    
    def _create_advanced_prompt(self, context: str, question: str) -> str:
        """Create advanced semantic-aware prompt (YOUR EXCELLENT ORIGINAL)"""
        return f"""You are an expert insurance policy analyst with advanced semantic understanding.

CONTEXT ANALYSIS FRAMEWORK:
- Apply deep semantic understanding to connect related concepts across documents
- Recognize implicit relationships and cross-references within policy content
- Understand hierarchical information structures and conditional dependencies
- Synthesize information from multiple sources with semantic coherence

DOCUMENT CONTEXT:
{context}

QUESTION: {question}

ADVANCED REASONING APPROACH:
1. SEMANTIC COMPREHENSION: Understand the full meaning and intent behind the question
2. CONTEXTUAL MAPPING: Map question elements to semantically relevant sections
3. RELATIONSHIP INFERENCE: Identify implicit connections between policy components
4. MULTI-SOURCE SYNTHESIS: Combine information while maintaining semantic consistency
5. CONDITIONAL REASONING: Apply logical reasoning to policy exceptions and conditions

RESPONSE REQUIREMENTS:
- Provide semantically rich, contextually grounded answers
- Include specific details: numbers, percentages, timeframes, conditions
- Write in clear, professional language without excessive quotes
- Address both explicit information and reasonable semantic inferences
- Structure information hierarchically when appropriate

ANSWER:"""
    
    def _clean_response(self, response: str) -> str:
        """Enhanced response cleaning (YOUR EXCELLENT ORIGINAL)"""
        # Remove excessive quotes
        response = re.sub(r'"([^"]{1,50})"', r'\1', response)
        response = re.sub(r'"(\w+)"', r'\1', response)
        response = re.sub(r'"(Rs\.?\s*[\d,]+[/-]*)"', r'\1', response)
        response = re.sub(r'"(\d+%)"', r'\1', response)
        response = re.sub(r'"(\d+\s*(?:days?|months?|years?))"', r'\1', response)
        
        # Clean policy references
        response = re.sub(r'[Aa]s stated in the policy[:\s]*"([^"]+)"', r'As per the policy, \1', response)
        response = re.sub(r'[Aa]ccording to the policy[:\s]*"([^"]+)"', r'According to the policy, \1', response)
        response = re.sub(r'[Tt]he policy states[:\s]*"([^"]+)"', r'The policy states that \1', response)
        
        # Fix spacing and formatting
        response = re.sub(r'\s+', ' ', response)
        response = response.replace(' ,', ',')
        response = response.replace(' .', '.')
        response = re.sub(r'\n\s*\n\s*\n+', '\n\n', response)
        
        return response.strip()

# --- AUTHENTICATION (YOUR EXCELLENT ORIGINAL) ---
async def verify_bearer_token(authorization: str = Header(None)):
    """Enhanced authentication with better logging"""
    if not authorization:
        raise HTTPException(status_code=401, detail="Authorization header required")
    
    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Invalid authorization format")
    
    token = authorization.replace("Bearer ", "")
    
    if len(token) < 10:
        raise HTTPException(status_code=401, detail="Invalid token format")
    
    logger.info(f"✅ Authentication successful with token: {token[:10]}...")
    return token

# --- GLOBAL INSTANCES (NO EARLY KAGGLE CONNECTION!) ---
multi_llm = MultiLLMManager()
doc_processor = UniversalDocumentProcessor()

# CRITICAL: Create lazy client (no immediate connection!)
kaggle_client = LazyKaggleModelClient()

# --- API MODELS ---
class SubmissionRequest(BaseModel):
    documents: List[str]
    questions: List[str]

class SubmissionResponse(BaseModel):
    answers: List[str]

# --- FIXED: BOTH GET AND POST ENDPOINTS FOR /api/v1/hackrx/run ---
@app.get("/api/v1/hackrx/run")
def test_endpoint():
    """GET endpoint for testing - fixes 405 Method Not Allowed error"""
    return {
        "message": "This endpoint requires POST method",
        "usage": "Send POST request with documents and questions",
        "status": "API is running - DEADLOCK-FREE with lazy initialization",
        "kaggle_connection": "Will initialize on first request",
        "fix": "Direct embedding management prevents async deadlocks",
        "method": "Use POST with JSON body",
        "example": {
            "documents": ["url1", "url2"],
            "questions": ["question1", "question2"]
        }
    }

# --- SPEED-OPTIMIZED MAIN ENDPOINT WITH GEMINI'S DEADLOCK FIX ---
@app.post("/api/v1/hackrx/run", response_model=SubmissionResponse, dependencies=[Depends(verify_bearer_token)])
async def run_submission(request: Request, submission_request: SubmissionRequest = Body(...)):
    start_time = time.time()
    logger.info(f"🎯 DEADLOCK-FREE KAGGLE-POWERED PROCESSING: {len(submission_request.documents)} docs, {len(submission_request.questions)} questions")
    
    try:
        # LAZY INITIALIZATION: Only now do we connect to Kaggle!
        logger.info("🔄 Initializing Kaggle connection (lazy initialization)...")
        
        # Check Kaggle health (this will trigger initialization)
        if not await kaggle_client.health_check():
            logger.error("❌ Kaggle endpoint not available!")
            return SubmissionResponse(answers=[
                "Model service unavailable" for _ in submission_request.questions
            ])
        
        # Create unique session with DEADLOCK-FREE pipeline
        session_id = f"kaggle_{uuid.uuid4().hex[:6]}"  # Shorter UUID
        rag_pipeline = DeadlockFreeRAGPipeline(session_id, multi_llm, kaggle_client)
        
        # Process all documents with higher concurrency
        all_chunks = []
        
        async with httpx.AsyncClient(
            timeout=45.0,
            headers={"ngrok-skip-browser-warning": "true"}
        ) as client:  # Tighter timeout + ngrok header
            # SPEED OPTIMIZATION: Higher concurrency
            semaphore = asyncio.Semaphore(5)  # Increased from 3
            
            async def process_single_document(doc_idx: int, doc_url: str):
                async with semaphore:
                    try:
                        logger.info(f"📥 Downloading document {doc_idx + 1}")
                        response = await client.get(doc_url, follow_redirects=True)
                        response.raise_for_status()
                        
                        # Get filename from URL or generate one
                        filename = os.path.basename(doc_url.split('?')[0]) or f"document_{doc_idx}"
                        
                        # Process document with caching
                        chunks = await doc_processor.process_document(filename, response.content)
                        
                        logger.info(f"✅ Document {doc_idx + 1}: {len(chunks)} chunks")
                        return chunks
                        
                    except Exception as e:
                        logger.error(f"❌ Document {doc_idx + 1} failed: {e}")
                        return []
            
            # Process all documents concurrently
            tasks = [
                process_single_document(i, url) 
                for i, url in enumerate(submission_request.documents)
            ]
            
            results = await asyncio.gather(*tasks)
            
            # Flatten results
            for chunks in results:
                all_chunks.extend(chunks)
        
        logger.info(f"📊 Total chunks processed: {len(all_chunks)}")
        
        if not all_chunks:
            logger.error("❌ No valid content extracted!")
            return SubmissionResponse(answers=[
                "No valid content could be extracted from the provided documents."
                for _ in submission_request.questions
            ])
        
        # Add to RAG pipeline with DEADLOCK-FREE processing
        await rag_pipeline.add_documents(all_chunks)
        
        # SPEED OPTIMIZATION: Full parallel question answering
        logger.info(f"⚡ Answering questions in parallel...")
        
        # INCREASED concurrency for questions
        semaphore = asyncio.Semaphore(4)  # Increased from 2
        
        async def answer_single_question(question: str) -> str:
            async with semaphore:
                return await rag_pipeline.answer_question(question)
        
        tasks = [answer_single_question(q) for q in submission_request.questions]
        answers = await asyncio.gather(*tasks)
        
        elapsed = time.time() - start_time
        logger.info(f"🎉 DEADLOCK-FREE KAGGLE-POWERED SUCCESS! Processed in {elapsed:.2f}s")
        
        return SubmissionResponse(answers=answers)
        
    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"💥 CRITICAL ERROR after {elapsed:.2f}s: {e}")
        
        return SubmissionResponse(answers=[
            "Processing error occurred. Please try again."
            for _ in submission_request.questions
        ])

# --- HEALTH ENDPOINTS (YOUR EXCELLENT ORIGINAL + DEADLOCK-FREE INFO) ---
@app.get("/")
def read_root():
    return {
        "message": "🎯 KAGGLE-POWERED HACKATHON RAG SYSTEM - DEADLOCK-FREE COMPLETE VERSION",
        "version": "5.4.0",
        "status": "FIXED: Deadlock-free + lazy initialization prevents all issues!",
        "target_time": "<20 seconds with Kaggle GPU",
        "supported_formats": list(doc_processor.processors.keys()),
        "features": [
            "Multi-format document processing (PDF, DOCX, Excel, CSV, HTML, etc.)",
            "Kaggle GPU-powered embeddings and reranking",
            "Multi-LLM fallback system (Groq, OpenAI, Gemini)",
            "Advanced semantic query enhancement",
            "Anti-jailbreak security system",
            "Optimized caching and concurrent processing",
            "Semantic chunking and context fusion",
            "R4 'half questions' handling",
            "Lightning-fast GPU-accelerated response times",
            "DEADLOCK-FREE async operations",
            "Lazy initialization prevents startup timeouts",
            "Direct embedding management"
        ],
        "kaggle_connection": "Lazy (connects on first API call)",
        "embedding_method": "Direct Kaggle management (no wrapper deadlock)",
        "fixes": [
            "DeadlockFreeRAGPipeline prevents async conflicts",
            "LazyKaggleModelClient prevents startup connection",
            "Direct embedding calls to Kaggle (no AsyncWrapper)",
            "Chroma as simple data store (no embedding function)",
            "CORS headers with ngrok-skip-browser-warning",
            "Both GET and POST endpoints for /api/v1/hackrx/run",
            "Improved error handling and logging",
            "Hugging Face Secrets support for dynamic URLs"
        ]
    }

@app.get("/health")
def health_check():
    return {
        "status": "healthy",
        "version": "5.4.0",
        "mode": "DEADLOCK_FREE_KAGGLE_GPU_POWERED_LAZY",
        "cache_size": len(doc_processor.cache),
        "kaggle_connection": "lazy (on-demand)",
        "embedding_method": "direct_kaggle_management",
        "timestamp": time.time(),
        "fixes_applied": [
            "deadlock_free_pipeline",
            "lazy_initialization",
            "direct_embedding_management", 
            "ngrok_compatibility",
            "http_method_fix",
            "cors_headers",
            "hf_secrets_support"
        ]
    }

@app.get("/test-kaggle")
async def test_kaggle_connection():
    """Test endpoint to check Kaggle connection (will trigger lazy initialization)"""
    try:
        is_healthy = await kaggle_client.health_check()
        return {
            "kaggle_connection": "initialized" if kaggle_client._initialized else "not_initialized",
            "health_status": "healthy" if is_healthy else "unhealthy",
            "endpoint": kaggle_client._endpoint if kaggle_client._initialized else "not_set",
            "timestamp": time.time()
        }
    except Exception as e:
        return {
            "kaggle_connection": "failed",
            "health_status": "error",
            "error": str(e),
            "timestamp": time.time()
        }

# --- RUN SERVER ---
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=7860)
